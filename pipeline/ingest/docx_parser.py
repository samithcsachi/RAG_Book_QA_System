from docx import Document
from pathlib import Path
from .parser_base import ParserBase
from typing import Tuple, Dict 

class DOCXParser(ParserBase):
    def extract_text_and_metadata(self, filepath: str) -> Tuple[str, Dict]:
        doc = Document(filepath)
        text_list = []
        for para in doc.paragraphs:
            text_list.append(para.text)
        text = "\n".join(text_list)
        metadata = {
            "filetype": "docx",
            "filename": str(Path(filepath).name),
            "num_paragraphs": len(doc.paragraphs)
        }
        return text, metadata